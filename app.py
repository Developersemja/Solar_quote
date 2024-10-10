from flask import Flask, request, send_file, render_template
from docx import Document
import io
import os

app = Flask(__name__)

WORD_TEMPLATE = 'profile.docx'  # Path to your Word template

@app.route('/')
def index():
    return render_template('form.html')

@app.route('/generate_quote', methods=['POST'])
def generate_quote():
    client_name = request.form['client_name']
    client_address = request.form['client_address']
    client_number = request.form['client_number']
    client_email = request.form['client_email']
    client_eb_no = request.form['client_eb_no']
    project_name = request.form['project_name']
    project_details = request.form['project_details']
    project_power_kw = request.form['project_power_kw']
    Project_Phase = request.form['Project_Phase']
    on_grid_off_grid = request.form['on_grid_off_grid']
    total_cost = request.form['total_cost']

    #------------------------table-------------------
    Solar_PV_module = request.form['Solar_PV_module']
    Solar_PV_module_no = request.form['Solar_PV_module_no']
    RCC_GI_Lot = request.form['RCC/GI:Lot']
    SPD_TYPE = request.form['SPD_TYPE']
    Solar_Inverter = request.form['Solar_Inverter']
    ACDB_with_SPD = request.form['ACDB_with_SPD']
    Cu_cable = request.form['Cu_cable']
    sqmm_Cu_cable = request.form['5sqmm_Cu_cable']
    copper_B = request.form['copper_B']
    Earth_cable = request.form['Earth_cable']
    LIGHTENING_ARRESTOR = request.form['LIGHTENING_ARRESTOR']
    Installation_kit = request.form['Installation_kit']
    MC_4_Cable = request.form['MC_4_Cable']

    # Generate an invoice number
    invoice_number = f"R{str(len(os.listdir()) + 1).zfill(3)}"

    doc = Document(WORD_TEMPLATE)
    
    # Replace placeholders with actual data
    replace_placeholder_in_runs(doc, '<<client_name>>', client_name)
    replace_placeholder_in_runs(doc, '<<client_address>>', client_address)
    replace_placeholder_in_runs(doc, '<<client_number>>', client_number)
    replace_placeholder_in_runs(doc, '<<client_email>>', client_email)
    replace_placeholder_in_runs(doc, '<<client_eb_no>>', client_eb_no)
    replace_placeholder_in_runs(doc, '<<project_name>>', project_name)
    replace_placeholder_in_runs(doc, '<<project_details>>', project_details)
    replace_placeholder_in_runs(doc, '<<project_power_kw>>', project_power_kw)
    replace_placeholder_in_runs(doc, '<<Project_Phase>>', Project_Phase)
    replace_placeholder_in_runs(doc, '<<on_grid_off_grid>>', on_grid_off_grid)
    replace_placeholder_in_runs(doc, '<<total_cost>>', total_cost)
    replace_placeholder_in_runs(doc, '<<invoice_number>>', invoice_number)

    '''--------------Table------------'''
    replace_placeholder_in_runs(doc, '<<Solar_PV_module>>', Solar_PV_module)
    replace_placeholder_in_runs(doc, '<<Solar_PV_module_no>>', Solar_PV_module_no)
    replace_placeholder_in_runs(doc, '<<RCC_GI_Lot>>', RCC_GI_Lot)
    replace_placeholder_in_runs(doc, '<<SPD_TYPE>>', SPD_TYPE)
    replace_placeholder_in_runs(doc, '<<Solar_Inverter>>', Solar_Inverter)
    replace_placeholder_in_runs(doc, '<<ACDB_with_SPD>>', ACDB_with_SPD)
    replace_placeholder_in_runs(doc, '<<Cu_cable>>', Cu_cable)
    replace_placeholder_in_runs(doc, '<<sqmm_Cu_cable>>', sqmm_Cu_cable)
    replace_placeholder_in_runs(doc, '<<copper_B>>', copper_B)
    replace_placeholder_in_runs(doc, '<<Earth_cable>>', Earth_cable)
    replace_placeholder_in_runs(doc, '<<LIGHTENING_ARRESTOR>>', LIGHTENING_ARRESTOR)
    replace_placeholder_in_runs(doc, '<<Installation_kit>>', Installation_kit)
    replace_placeholder_in_runs(doc, '<<MC_4_Cable>>', MC_4_Cable)
    
    # Save the updated document in memory using io.BytesIO
    memory_file = io.BytesIO()
    doc.save(memory_file)
    memory_file.seek(0)

    # Return the generated Word document as an attachment
    return send_file(memory_file, as_attachment=True, download_name=f"quote_{invoice_number}_{client_name.replace(' ', '_')}.docx")

def replace_placeholder_in_runs(doc, placeholder, replacement_text):
    """Replace placeholders in runs while preserving formatting, both in normal paragraphs and tables."""
    
    # Replace in normal paragraphs
    for paragraph in doc.paragraphs:
        full_text = ''.join([run.text for run in paragraph.runs])
        if placeholder in full_text:
            # Replace the placeholder in the full text
            full_text = full_text.replace(placeholder, replacement_text)
            
            # Clear all the runs in the paragraph
            for run in paragraph.runs:
                run.text = ''
            
            # Put the replaced text back in the first run
            paragraph.runs[0].text = full_text
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join([run.text for run in paragraph.runs])
                    if placeholder in full_text:
                        # Replace the placeholder in the full text
                        full_text = full_text.replace(placeholder, replacement_text)
                        
                        # Clear all the runs in the paragraph
                        for run in paragraph.runs:
                            run.text = ''
                        
                        # Put the replaced text back in the first run
                        paragraph.runs[0].text = full_text

if __name__ == '__main__':
    app.run(debug=True)
